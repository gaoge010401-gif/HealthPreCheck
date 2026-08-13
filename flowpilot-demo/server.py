from __future__ import annotations

import cgi
import io
import json
import shutil
import subprocess
import tempfile
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path


ROOT = Path(__file__).resolve().parent
MAX_UPLOAD_BYTES = 10 * 1024 * 1024


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(data: bytes) -> str:
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()
    except Exception:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def extract_image(data: bytes, suffix: str) -> tuple[str, str]:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return (
            "",
            "Image OCR needs Tesseract.js in the browser or a server with Tesseract installed.",
        )

    with tempfile.NamedTemporaryFile(suffix=suffix, delete=True) as image_file:
        image_file.write(data)
        image_file.flush()
        result = subprocess.run(
            [tesseract, image_file.name, "stdout", "-l", "eng"],
            check=False,
            capture_output=True,
            text=True,
        )
    if result.returncode != 0:
        return "", result.stderr.strip() or "Tesseract OCR failed."
    return result.stdout.strip(), ""


def extract_text(filename: str, content_type: str, data: bytes) -> dict:
    suffix = Path(filename).suffix.lower()
    warning = ""
    method = "server text extraction"

    if suffix in {".txt", ".csv"} or content_type.startswith("text/"):
        text = decode_text(data)
        method = "server text reader"
    elif suffix == ".docx":
        text = extract_docx(data)
        method = "server DOCX extractor"
    elif suffix == ".pdf":
        text = extract_pdf(data)
        method = "server PDF text extractor"
        if not text:
            warning = "PDF did not contain selectable text. Try uploading a screenshot/image for OCR."
    elif content_type.startswith("image/") or suffix in {".png", ".jpg", ".jpeg"}:
        text, warning = extract_image(data, suffix or ".png")
        method = "server OCR"
    else:
        text = ""
        warning = "Unsupported file type for this prototype."

    return {
        "filename": filename,
        "method": method,
        "text": text,
        "warning": warning,
    }


class HealthPreCheckHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(ROOT), **kwargs)

    def end_headers(self) -> None:
        self.send_header("Cache-Control", "no-store")
        super().end_headers()

    def send_json(self, status: int, payload: dict) -> None:
        data = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def do_POST(self) -> None:
        if self.path != "/api/extract":
            self.send_error(404)
            return

        content_length = int(self.headers.get("Content-Length", "0"))
        if content_length > MAX_UPLOAD_BYTES:
            self.send_json(413, {"error": "File too large for prototype upload."})
            return

        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                "CONTENT_LENGTH": str(content_length),
            },
        )
        file_item = form["document"] if "document" in form else None
        if file_item is None or not getattr(file_item, "file", None):
            self.send_json(400, {"error": "Missing uploaded document."})
            return

        filename = Path(file_item.filename or "uploaded-document").name
        data = file_item.file.read()
        try:
            payload = extract_text(filename, file_item.type or "", data)
            self.send_json(200, payload)
        except Exception as error:
            self.send_json(500, {"error": str(error)})


def main() -> None:
    server = ThreadingHTTPServer(("127.0.0.1", 4173), HealthPreCheckHandler)
    print("HealthPreCheck demo server running at http://127.0.0.1:4173")
    server.serve_forever()


if __name__ == "__main__":
    main()
