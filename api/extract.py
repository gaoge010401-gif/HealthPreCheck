from __future__ import annotations

import cgi
import io
import json
from http.server import BaseHTTPRequestHandler
from pathlib import Path
from urllib.parse import urlparse


MAX_UPLOAD_BYTES = 10 * 1024 * 1024
STATIC_ROOT = Path(__file__).resolve().parents[1] / "flowpilot-demo"
STATIC_ROUTES = {
    "/": "index.html",
    "/index.html": "index.html",
    "/app.js": "app.js",
    "/styles.css": "styles.css",
    "/flowpilot-demo/index.html": "index.html",
    "/flowpilot-demo/app.js": "app.js",
    "/flowpilot-demo/styles.css": "styles.css",
}
CONTENT_TYPES = {
    ".html": "text/html; charset=utf-8",
    ".css": "text/css; charset=utf-8",
    ".js": "application/javascript; charset=utf-8",
}


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages).strip()


def extract_text(filename: str, content_type: str, data: bytes) -> dict:
    suffix = Path(filename).suffix.lower()
    warning = ""
    method = "server text extraction"

    if suffix in {".txt", ".csv"} or content_type.startswith("text/"):
        text = decode_text(data)
        method = "server text reader"
    elif suffix == ".docx":
        text = extract_docx(data)
        method = "server DOCX extractor"
    elif suffix == ".pdf":
        text = extract_pdf(data)
        method = "server PDF text extractor"
        if not text:
            warning = "PDF did not contain selectable text. Try browser image OCR or paste text manually."
    elif content_type.startswith("image/") or suffix in {".png", ".jpg", ".jpeg"}:
        text = ""
        warning = "Server-side OCR is not enabled on Vercel. Browser OCR will run first when Tesseract.js is available."
        method = "server OCR fallback"
    else:
        text = ""
        warning = "Unsupported file type for this prototype."

    return {
        "filename": filename,
        "method": method,
        "text": text,
        "warning": warning,
    }


class handler(BaseHTTPRequestHandler):
    def send_json(self, status: int, payload: dict) -> None:
        data = json.dumps(payload).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        self.wfile.write(data)

    def send_file(self, status: int, path: Path, include_body: bool = True) -> None:
        data = path.read_bytes()
        self.send_response(status)
        self.send_header("Content-Type", CONTENT_TYPES.get(path.suffix, "application/octet-stream"))
        self.send_header("Content-Length", str(len(data)))
        self.end_headers()
        if include_body:
            self.wfile.write(data)

    def serve_static(self, include_body: bool = True) -> None:
        request_path = urlparse(self.path).path
        filename = STATIC_ROUTES.get(request_path)
        if not filename:
            self.send_json(404, {"error": "Not found"})
            return

        path = STATIC_ROOT / filename
        if not path.exists():
            self.send_json(404, {"error": "Static asset missing"})
            return

        self.send_file(200, path, include_body)

    def do_GET(self) -> None:
        self.serve_static()

    def do_HEAD(self) -> None:
        self.serve_static(include_body=False)

    def do_POST(self) -> None:
        content_length = int(self.headers.get("Content-Length", "0"))
        if content_length <= 0:
            self.send_json(400, {"error": "Missing uploaded document."})
            return
        if content_length > MAX_UPLOAD_BYTES:
            self.send_json(413, {"error": "File too large for prototype upload."})
            return
        if "multipart/form-data" not in self.headers.get("Content-Type", ""):
            self.send_json(400, {"error": "Expected multipart form upload."})
            return

        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                "CONTENT_LENGTH": str(content_length),
            },
        )
        file_item = form["document"] if "document" in form else None
        if file_item is None or not getattr(file_item, "file", None):
            self.send_json(400, {"error": "Missing uploaded document."})
            return

        filename = Path(file_item.filename or "uploaded-document").name
        data = file_item.file.read()
        try:
            self.send_json(200, extract_text(filename, file_item.type or "", data))
        except Exception as error:
            self.send_json(500, {"error": str(error)})
